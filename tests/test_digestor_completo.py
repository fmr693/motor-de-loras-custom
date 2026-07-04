"""
test_digestor_completo.py
=========================
Prueba COMPLETA del DataDigestor v3: TODOS los formatos + export universal.
"""

import sys, os, tempfile, json, wave, struct, math
from pathlib import Path

import pytest
from motor.digestor import DataDigestor, detect_file_type

PASS = FAIL = 0
def check(cond, msg):
    global PASS, FAIL
    if cond: PASS += 1; print(f"  ✅ {msg}")
    else: FAIL += 1; print(f"  ❌ {msg}")

def section(title):
    print(f"\n{'='*60}\n  {title}\n{'='*60}")

TMP = Path(tempfile.mkdtemp(prefix="digestor_full_"))
OUT = Path(__file__).resolve().parent / "test_digestor_completo_output"
OUT.mkdir(exist_ok=True)
import pandas as pd

# titanic.csv fue eliminado del repo. Saltar el módulo si no está presente.
_TITANIC_CSV = Path(__file__).resolve().parent / "titanic.csv"
if not _TITANIC_CSV.exists():
    pytest.skip(
        reason="titanic.csv no encontrado en tests/ — cópialo para ejecutar estas pruebas.",
        allow_module_level=True,
    )

# ═══════════════════════════════════════════════════════════════════
section("F1 — detect_file_type (todos los formatos)")
# ═══════════════════════════════════════════════════════════════════

expected = {".csv":"csv", ".json":"json", ".txt":"txt", ".pdf":"pdf",
            ".docx":"docx", ".html":"html", ".xlsx":"excel",
            ".mp3":"audio", ".wav":"audio", ".mp4":"video", ".mkv":"video",
            ".png":"image", ".jpg":"image"}
for ext, expect in expected.items():
    result = detect_file_type(f"test{ext}")
    check(result == expect, f"{ext} → {result} (esperado: {expect})")

# ═══════════════════════════════════════════════════════════════════
section("F2 — from_csv() + semáforo + task variation")
# ═══════════════════════════════════════════════════════════════════

d = DataDigestor(
    task="¿Sobrevivió este pasajero al Titanic? Responde YES o NO.",
    label_col="Survived", label_map={0:"NO", 1:"YES"},
    auto_enrich=False,
)
d.from_csv(str(Path(__file__).resolve().parent / "titanic.csv"))
check(len(d._examples) == 891, f"CSV: {len(d._examples)} ejemplos (esperado 891)")
result = d.validate(verbose=False)
check(result["total"] == 891, f"validate: {result['total']} total")
check(result["semaforo"] == "VERDE", f"Semáforo: {result['semaforo']} (esperado VERDE)")

# Verificar task variation
tasks = set()
for ex in d._examples[:50]:
    msgs = ex.get("messages", [])
    user = next((m["content"] for m in msgs if m["role"]=="user"), "")
    task_part = user.split("\n\n")[0]
    tasks.add(task_part)
check(len(tasks) >= 2, f"Task variation activa: {len(tasks)} variantes distintas")

# ═══════════════════════════════════════════════════════════════════
section("F3 — from_txt() (texto tabulado)")
# ═══════════════════════════════════════════════════════════════════

sms_txt = str(TMP / "sms.txt")
with open(sms_txt, "w", encoding="utf-8") as f:
    f.write("ham\tHey, how are you?\nspam\tWIN A FREE IPHONE NOW!!!\nham\tSee you at 7\nspam\tURGENT: verify account")

d2 = DataDigestor(task="¿Es SPAM o HAM?", label_col="label",
                  label_map={"spam":"SPAM","ham":"HAM"}, auto_enrich=False)
d2.from_txt(sms_txt, delimiter="\t", text_col_idx=1, label_col_idx=0)
check(len(d2._examples) == 4, f"TXT tabulado: {len(d2._examples)} ejemplos")

# ═══════════════════════════════════════════════════════════════════
section("F4 — from_json() (array de objetos)")
# ═══════════════════════════════════════════════════════════════════

json_path = str(TMP / "data.json")
with open(json_path, "w", encoding="utf-8") as f:
    json.dump([
        {"texto": "Paciente con diabetes tipo 2, HbA1c 8.2%.", "clase": "medico"},
        {"texto": "Las acciones de Apple suben 5% tras resultados.", "clase": "financiero"},
        {"texto": "NullPointerException en UserService.java:342.", "clase": "tecnico"},
    ], f)

d3 = DataDigestor(task="Clasifica el documento.", label_col="clase", auto_enrich=False)
d3.from_json(json_path, text_field="texto", label_field="clase")
check(len(d3._examples) == 3, f"JSON: {len(d3._examples)} ejemplos")

# ═══════════════════════════════════════════════════════════════════
section("F5 — from_docx()")
# ═══════════════════════════════════════════════════════════════════

from docx import Document
doc = Document()
doc.add_paragraph("INFORME MEDICO: Paciente femenina 67 anios con diabetes.")
doc.add_paragraph("Tratamiento: metformina 850mg cada 12 horas.")
doc.add_paragraph("El control glucemico muestra mejoria significativa.")
docx_path = str(TMP / "medico.docx")
doc.save(docx_path)

d4 = DataDigestor(task="Resume este documento.", auto_enrich=True)
d4.from_docx(docx_path)
check(len(d4._examples) >= 3, f"DOCX: {len(d4._examples)} parrafos")
check(any("medico" in m["content"].lower() or "medical" in m["content"].lower()
          for ex in d4._examples for m in ex.get("messages",[]) if m["role"]=="user"),
      "Contenido medico presente en DOCX")

# ═══════════════════════════════════════════════════════════════════
section("F6 — from_html()")
# ═══════════════════════════════════════════════════════════════════

html = """<html><body><script>alert('xss')</script><style>.ad{}</style>
<article><h1>NASDAQ records</h1><p>El NASDAQ sube 2.3% liderado por Tesla.</p>
<p>Goldman Sachs: BULLISH. EBITDA supera expectativas.</p></article></body></html>"""
html_path = str(TMP / "finanzas.html")
with open(html_path, "w", encoding="utf-8") as f:
    f.write(html)

d5 = DataDigestor(task="Clasifica esta noticia.", auto_enrich=False)
d5.from_html(html_path, text_selector="article")
check(len(d5._examples) >= 1, f"HTML: {len(d5._examples)} bloques")
check(not any("alert" in m["content"] for ex in d5._examples for m in ex.get("messages",[]) if m["role"]=="user"),
      "JavaScript no aparece en salida HTML")

# ═══════════════════════════════════════════════════════════════════
section("F7 — from_pdf()")
# ═══════════════════════════════════════════════════════════════════

try:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "CONTRATO DE ARRENDAMIENTO\n\n"
                   "Arrendador: Carlos Lopez. Arrendataria: Maria Fernandez.\n"
                   "Renta: 950 EUR/mes. Duracion: 12 meses prorrogables.\n"
                   "El incumplimiento dara lugar a rescision del contrato.")
    pdf_path = str(TMP / "contrato.pdf")
    pdf.output(pdf_path)

    d6 = DataDigestor(task="Clasifica este documento legal.", auto_enrich=True)
    d6.from_pdf(pdf_path)
    check(len(d6._examples) >= 1, f"PDF: {len(d6._examples)} paginas")
except ImportError:
    print("  ⚠️ fpdf2 no instalado — saltando test PDF")

# ═══════════════════════════════════════════════════════════════════
section("F8 — from_audio()")
# ═══════════════════════════════════════════════════════════════════

try:
    from faster_whisper import WhisperModel
    # Generar audio sintetico de 3 segundos
    wav_path = str(TMP / "test.wav")
    w = wave.open(wav_path, 'w')
    w.setnchannels(1); w.setsampwidth(2); w.setframerate(16000)
    for i in range(16000 * 3):
        w.writeframes(struct.pack('<h', int(8000 * math.sin(2*math.pi*440*i/16000))))
    w.close()

    d7 = DataDigestor(task="Transcripcion de audio.", auto_enrich=False)
    d7.from_audio(wav_path)
    # El audio sintetico no tiene voz → puede dar 0 fragmentos o pocos
    check(True, f"from_audio() ejecutado sin errores ({len(d7._examples)} fragmentos)")
except ImportError:
    print("  ⚠️ faster-whisper o modelo no disponible — saltando test audio")

# ═══════════════════════════════════════════════════════════════════
section("F9 — from_video() (si ffmpeg disponible)")
# ═══════════════════════════════════════════════════════════════════

import subprocess
try:
    subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
    # Generar video minimo de 1 segundo
    video_path = str(TMP / "test.mp4")
    subprocess.run([
        "ffmpeg", "-f", "lavfi", "-i", "color=c=black:s=64x64:d=1",
        "-f", "lavfi", "-i", "sine=f=440:d=1",
        "-c:v", "libx264", "-c:a", "aac", "-shortest",
        "-y", video_path
    ], capture_output=True, check=True)

    d8 = DataDigestor(task="Transcripcion de video.", auto_enrich=False)
    d8.from_video(video_path)
    check(True, f"from_video() ejecutado sin errores ({len(d8._examples)} fragmentos)")
except (FileNotFoundError, subprocess.CalledProcessError):
    print("  ⚠️ ffmpeg no disponible — saltando test video")

# ═══════════════════════════════════════════════════════════════════
section("F10 — Export universal (G2)")
# ═══════════════════════════════════════════════════════════════════

# Usar Titanic dataset ya cargado
n_u = d.to_unsloth(str(OUT / "titanic_unsloth.jsonl"))
with open(OUT / "titanic_unsloth.jsonl", encoding="utf-8") as f:
    first_u = json.loads(f.readline())
check("instruction" in first_u and "output" in first_u, f"Unsloth: formato Alpaca correcto ({n_u} ejemplos)")

n_l = d.to_llamafactory(str(OUT / "llamafactory"), "titanic")
check((OUT / "llamafactory" / "dataset_info.json").exists(), f"LLaMA-Factory: dataset_info.json generado ({n_l} ejemplos)")

n_a = d.to_axolotl(str(OUT / "axolotl"), "titanic")
check((OUT / "axolotl" / "axolotl_config.yml").exists(), f"Axolotl: YAML generado ({n_a} ejemplos)")

# ═══════════════════════════════════════════════════════════════════
section("F11 — Detección de dominio")
# ═══════════════════════════════════════════════════════════════════

# PDF legal deberia detectar dominio legal
if 'd6' in dir() and d6._examples:
    check(d6._detected_domain is not None, f"Dominio detectado: {d6._detected_domain}")

# CSV medico (no tenemos etiquetas reales, pero from_docx con texto medico)
check(d4._detected_domain is not None, f"Dominio DOCX: {d4._detected_domain}")

# ═══════════════════════════════════════════════════════════════════
section("F12 — Semáforo (todos los escenarios)")
# ═══════════════════════════════════════════════════════════════════

# ROJO: pocos ejemplos
d_tiny = DataDigestor(task="test", label_col="label", auto_enrich=False)
d_tiny._examples = [{"messages": [{"role":"user","content":"a"},{"role":"assistant","content":"b"}]}] * 10
r_tiny = d_tiny.validate(verbose=False)
check(r_tiny["semaforo"] == "ROJO", f"Semáforo ROJO (10 ejemplos): {r_tiny['semaforo']}")

# VERDE: muchos ejemplos
check(result["semaforo"] == "VERDE", f"Semáforo VERDE (891 ejemplos): {result['semaforo']}")

# Tarea homogenea detectada
warnings_str = " ".join(r_tiny.get("warnings", []))
check("HOMOG" in warnings_str.upper(), "Tarea homogenea detectada en dataset pequeno")

# ═══════════════════════════════════════════════════════════════════
section(f"RESULTADO FINAL: {PASS} PASS / {FAIL} FAIL")
print(f"  Output: {OUT}")
if FAIL == 0:
    print(f"\n  🎉 TODOS LOS TESTS PASARON — DataDigestor v3 OPERATIVO")
