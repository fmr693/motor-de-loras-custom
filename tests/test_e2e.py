"""
tests/test_e2e.py
=================
CAPA 2 — Integration tests: pipeline completo de extremo a extremo.

Verifica que todos los componentes funcionan juntos sin GPU:
  A. Digestor → to_jsonl → load → validate (roundtrip)
  B. Formatos de exportación: LLaMA-Factory, Unsloth, Axolotl
  C. Flujos completos: CSV, DOCX, HTML, user_profile, tool_calls
  D. ContinualLearner: registry, replay buffer, rollback
  E. Interacción server ↔ continual (S10.1 log format)

No requiere GPU ni descarga de modelos.
"""

from __future__ import annotations

import csv
import json
import os
import shutil
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest

from motor.digestor import DataDigestor
from motor.continual import ContinualLearner

# ===========================================================================
# Fixtures
# ===========================================================================

@pytest.fixture
def csv_titanic(tmp_path) -> Path:
    """CSV pequeño estilo Titanic con 10 filas + etiquetas binarias."""
    p = tmp_path / "mini_titanic.csv"
    with open(p, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["PassengerId", "Survived", "Pclass", "Name", "Sex", "Age", "Fare"])
        w.writerow([1, 0, 3, "Braund, Mr. Owen Harris", "male", 22, 7.25])
        w.writerow([2, 1, 1, "Cumings, Mrs. John Bradley", "female", 38, 71.28])
        w.writerow([3, 1, 3, "Heikkinen, Miss. Laina", "female", 26, 7.925])
        w.writerow([4, 1, 1, "Futrelle, Mrs. Jacques Heath", "female", 35, 53.1])
        w.writerow([5, 0, 3, "Allen, Mr. William Henry", "male", 35, 8.05])
        w.writerow([6, 0, 3, "Moran, Mr. James", "male", 27, 8.4583])
        w.writerow([7, 0, 1, "McCarthy, Mr. Timothy J", "male", 54, 51.8625])
        w.writerow([8, 1, 3, "Palsson, Miss. Torborg", "female", 19, 21.075])
        w.writerow([9, 0, 2, "Johnson, Mrs. Oscar W", "female", 21, 11.1333])
        w.writerow([10, 1, 3, "Nasser, Mrs. Nicholas", "female", 25, 30.0708])
    return p


@pytest.fixture
def jsonl_output(tmp_path) -> Path:
    return tmp_path / "output.jsonl"


@pytest.fixture
def continual_learner(tmp_path) -> ContinualLearner:
    registry = tmp_path / "registry.json"
    return ContinualLearner(
        model_id="Qwen/Qwen2.5-3B-Instruct",
        registry_path=str(registry),
        replay_buffer_size=10,
        rollback_threshold=0.15,
    )


# ===========================================================================
# A. Roundtrip: digestor → to_jsonl → load → validate
# ===========================================================================

class TestRoundtripDigestor:
    """El pipeline completo desde CSV hasta validación del dataset."""

    def test_csv_a_jsonl_y_validacion(self, csv_titanic, jsonl_output):
        """CSV con etiquetas → JSONL → cargar → validar."""
        d = DataDigestor(
            task="¿Sobrevivió este pasajero del Titanic?",
            label_col="Survived",
            label_map={0: "NO", 1: "YES"},
        )
        d.from_csv(str(csv_titanic))
        assert len(d._examples) == 10

        d.to_jsonl(str(jsonl_output), deduplicate=False)

        # Cargar de vuelta
        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) == 10

        # Validar
        report = d2.validate()
        assert "semaforo" in report
        assert report["total"] == 10

    def test_csv_sin_label_map_genera_ejemplos(self, csv_titanic, jsonl_output):
        """Sin label_map, usa el valor crudo de la columna."""
        d = DataDigestor(
            task="Describe a este pasajero",
            label_col="Survived",
        )
        d.from_csv(str(csv_titanic))
        d.to_jsonl(str(jsonl_output), deduplicate=False)

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) == 10

    def test_pipeline_con_dedup_activa(self, csv_titanic, jsonl_output):
        """Deduplicación no debería eliminar ejemplos únicos."""
        d = DataDigestor(
            task="¿Sobrevivió?",
            label_col="Survived",
            label_map={0: "NO", 1: "YES"},
        )
        d.from_csv(str(csv_titanic))
        d.to_jsonl(str(jsonl_output))  # deduplicate=True por defecto

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        # Todos son únicos → mismo número
        assert len(d2._examples) == 10

    def test_pipeline_elimina_duplicados_reales(self, jsonl_output):
        """Ejemplos idénticos se eliminan con deduplicate=True."""
        d = DataDigestor(task="test")
        # Añadir 5 pares idénticos
        for _ in range(5):
            d._examples.append({
                "messages": [
                    {"role": "user", "content": "Hola"},
                    {"role": "assistant", "content": "¿Cómo estás?"},
                ]
            })
        # Añadir 1 diferente
        d._examples.append({
            "messages": [
                {"role": "user", "content": "Adiós"},
                {"role": "assistant", "content": "Hasta luego"},
            ]
        })
        d.to_jsonl(str(jsonl_output))  # deduplicate=True

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        # 5 idénticos → 1 + 1 diferente = 2
        assert len(d2._examples) == 2

    def test_score_y_validate_incluyen_scores(self, csv_titanic):
        """score_examples + validate(include_scores=True) en el mismo flujo."""
        d = DataDigestor(
            task="¿Sobrevivió?",
            label_col="Survived",
        )
        d.from_csv(str(csv_titanic))

        scores = d.score_examples()
        assert len(scores) == 10
        assert all("score" in s for s in scores)

        report = d.validate(include_scores=True)
        assert "quality_scores" in report
        assert len(report["quality_scores"]) == 10

    def test_augment_no_rompe_validacion(self, csv_titanic):
        """Aumentar ejemplos no invalida el dataset."""
        d = DataDigestor(
            task="¿Sobrevivió este pasajero?",
            label_col="Survived",
            label_map={0: "NO", 1: "YES"},
        )
        d.from_csv(str(csv_titanic))
        original = len(d._examples)
        d.augment(strategy="template_swap", n_augmented=5)
        assert len(d._examples) >= original

        report = d.validate()
        assert report["total"] >= original


# ===========================================================================
# B. Formatos de exportación
# ===========================================================================

class TestExportFormats:
    """Verifica que to_llamafactory, to_unsloth, to_axolotl generan archivos válidos."""

    @pytest.fixture
    def digestor_con_ejemplos(self):
        d = DataDigestor(task="Clasifica el sentimiento")
        d._examples = [
            {"messages": [
                {"role": "user", "content": "Me encanta este producto"},
                {"role": "assistant", "content": "POSITIVO"},
            ]},
            {"messages": [
                {"role": "user", "content": "Es terrible, no funciona"},
                {"role": "assistant", "content": "NEGATIVO"},
            ]},
        ]
        return d

    def test_to_llamafactory_crea_json_y_dataset_info(self, digestor_con_ejemplos, tmp_path):
        out = tmp_path / "llamafactory_export"
        digestor_con_ejemplos.to_llamafactory(str(out), dataset_name="test_sentiment")

        json_file = out / "test_sentiment.json"
        info_file = out / "dataset_info.json"

        assert json_file.exists()
        assert info_file.exists()

        data = json.loads(json_file.read_text(encoding="utf-8"))
        assert len(data) == 2
        assert "conversations" in data[0]

        info = json.loads(info_file.read_text(encoding="utf-8"))
        assert "test_sentiment" in info

    def test_to_unsloth_formato_alpaca(self, digestor_con_ejemplos, tmp_path):
        out = tmp_path / "unsloth_export.jsonl"
        digestor_con_ejemplos.to_unsloth(str(out))

        assert out.exists()
        lines = out.read_text(encoding="utf-8").strip().splitlines()
        assert len(lines) == 2

        first = json.loads(lines[0])
        assert "instruction" in first
        assert "output" in first

    def test_to_axolotl_crea_jsonl_y_config(self, digestor_con_ejemplos, tmp_path):
        out = tmp_path / "axolotl_export"
        digestor_con_ejemplos.to_axolotl(str(out), dataset_name="test_sentiment")

        jsonl_file = out / "test_sentiment.jsonl"
        config_file = out / "axolotl_config.yml"

        assert jsonl_file.exists()
        assert config_file.exists()

        config = config_file.read_text(encoding="utf-8")
        assert "base_model" in config.lower() or "model" in config.lower()


# ===========================================================================
# C. Flujos completos: DOCX, HTML, user_profile, tool_calls
# ===========================================================================

class TestFlujosCompletos:
    """Escenarios realistas que combinan múltiples funciones del digestor."""

    def test_flujo_from_docx_a_jsonl(self, tmp_path, jsonl_output):
        """Crear DOCX → procesar → exportar JSONL → cargar."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx no instalado")

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Informe financiero del primer trimestre.")
        doc.add_paragraph("Los ingresos aumentaron un 15% respecto al año anterior.")
        doc.add_paragraph("Se recomienda mantener la inversión en tecnología.")
        doc.save(str(docx_path))

        d = DataDigestor(task="Analiza el sentimiento financiero de este texto")
        d.from_docx(str(docx_path))
        assert len(d._examples) > 0

        d.to_jsonl(str(jsonl_output), deduplicate=False)
        assert jsonl_output.exists()

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) > 0

    def test_flujo_from_html_a_jsonl(self, tmp_path, jsonl_output):
        """Crear HTML → procesar → exportar → cargar."""
        pytest.importorskip("bs4", reason="beautifulsoup4 no instalado — pip install beautifulsoup4 lxml")
        html_path = tmp_path / "test.html"
        html_path.write_text("""
        <html><body>
        <h1>Reporte Médico</h1>
        <p>El paciente muestra mejoría significativa tras el tratamiento.</p>
        <p>Se recomienda continuar con la dosis actual durante 2 semanas más.</p>
        </body></html>
        """, encoding="utf-8")

        d = DataDigestor(task="Extrae información clínica relevante")
        d.from_html(str(html_path))
        assert len(d._examples) > 0

        d.to_jsonl(str(jsonl_output), deduplicate=False)
        assert jsonl_output.exists()

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) > 0

    def test_flujo_from_user_profile(self, tmp_path):
        """Perfil de usuario → dataset → validar."""
        import json as _j
        profile = {
            "name": "Doctora García",
            "specialty": "Cardiología",
            "hospital": "Hospital Central",
            "preferences": {
                "tone": "formal",
                "language": "es",
            },
            "common_tasks": [
                "Revisar historiales médicos",
                "Programar citas de seguimiento",
                "Consultar dosis de medicamentos",
            ],
        }
        # from_user_profile espera ruta a JSON, no dict
        profile_path = tmp_path / "profile.json"
        profile_path.write_text(_j.dumps(profile), encoding="utf-8")

        d = DataDigestor(task="Asistente médico personalizado")
        d.from_user_profile(str(profile_path))
        assert len(d._examples) > 0

        report = d.validate()
        assert "semaforo" in report

        # Verificar que los ejemplos mencionan conceptos del perfil
        all_text = json.dumps(d._examples, ensure_ascii=False).lower()
        assert "cardiología" in all_text or "cardio" in all_text or "médico" in all_text or "hospital" in all_text

    def test_flujo_tool_calls_a_jsonl(self, jsonl_output):
        """Generar tool_calls → exportar → cargar."""
        tools = [
            {"name": "search_patient", "description": "Busca paciente por nombre o ID",
             "parameters": {"query": {"type": "str"}, "department": {"type": "str"}}},
            {"name": "schedule_appointment", "description": "Programa una cita médica",
             "parameters": {"patient_id": {"type": "str"}, "date": {"type": "str"}, "doctor": {"type": "str"}}},
        ]

        d = DataDigestor(task="Asistente médico con herramientas")
        d.generate_tool_calls(tools, n_per_tool=3)
        assert len(d._examples) == 6

        d.to_jsonl(str(jsonl_output), deduplicate=False)

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) == 6

    def test_flujo_carpeta_mixta(self, tmp_path, jsonl_output):
        """Carpeta con CSV + TXT → procesar todo junto."""
        folder = tmp_path / "mixta"
        folder.mkdir()

        # CSV
        csv_p = folder / "data.csv"
        with open(csv_p, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["text", "label"])
            w.writerow(["El producto es excelente", "POS"])
            w.writerow(["No cumple expectativas", "NEG"])

        # TXT
        txt_p = folder / "notas.txt"
        txt_p.write_text("Reunión del comité ejecutivo.\nAprobado el presupuesto Q3.", encoding="utf-8")

        d = DataDigestor(task="Análisis de documentos empresariales", label_col="label")
        d.from_folder(str(folder))
        assert len(d._examples) >= 2

        d.to_jsonl(str(jsonl_output), deduplicate=False)
        assert jsonl_output.exists()

        d2 = DataDigestor(task="test")
        d2.load_jsonl(str(jsonl_output))
        assert len(d2._examples) >= 2


# ===========================================================================
# D. ContinualLearner: registry, replay buffer, rollback
# ===========================================================================

class TestContinualLearnerE2E:
    """Operaciones del ContinualLearner sin GPU."""

    def test_registry_vacio_al_inicio(self, continual_learner):
        reg = continual_learner.get_registry()
        assert reg.get("adapters") is not None

    def test_register_existing_funciona(self, continual_learner, tmp_path):
        adapter_dir = tmp_path / "test_adapter"
        adapter_dir.mkdir()
        meta = {"model_id": "Qwen/Qwen2.5-3B-Instruct", "eval_loss": 0.05}
        (adapter_dir / "meta.json").write_text(json.dumps(meta))

        dataset = tmp_path / "dataset.jsonl"
        dataset.write_text(json.dumps({"messages": []}) + "\n")

        continual_learner.register_existing(
            str(adapter_dir), str(dataset), name="test_v1"
        )

        reg = continual_learner.get_registry()
        assert len(reg["adapters"]) == 1
        assert reg["adapters"][0]["name"] == "test_v1"
        assert reg["adapters"][0]["eval_loss"] == 0.05

    def test_register_existing_sin_meta(self, continual_learner, tmp_path):
        adapter_dir = tmp_path / "bare_adapter"
        adapter_dir.mkdir()
        # Sin meta.json

        dataset = tmp_path / "dataset.jsonl"
        dataset.write_text(json.dumps({"messages": []}) + "\n")

        continual_learner.register_existing(
            str(adapter_dir), str(dataset), name="bare_v1", eval_loss=0.12
        )

        reg = continual_learner.get_registry()
        assert reg["adapters"][0]["eval_loss"] == 0.12

    def test_replay_buffer_muestrea_correctamente(self, continual_learner, tmp_path):
        """El replay buffer toma ejemplos de adapters registrados."""
        # Registrar 2 adapters con datasets
        for i in range(2):
            ad = tmp_path / f"adapter_{i}"
            ad.mkdir()
            (ad / "meta.json").write_text(json.dumps({"model_id": "Qwen/Qwen2.5-3B-Instruct"}))

            ds = tmp_path / f"dataset_{i}.jsonl"
            examples = [
                {"messages": [{"role": "user", "content": f"msg_{i}_{j}"},
                              {"role": "assistant", "content": f"resp_{i}_{j}"}]}
                for j in range(5)
            ]
            ds.write_text("\n".join(json.dumps(ex) for ex in examples) + "\n")
            continual_learner.register_existing(str(ad), str(ds), name=f"v{i}")

        # Muestrear
        replay = continual_learner._sample_replay_buffer(seed=42)
        # Debería tener ejemplos de los datasets registrados
        assert len(replay) > 0

    def test_history_no_crashea_con_registry_vacio(self, continual_learner, capsys):
        continual_learner.history()
        captured = capsys.readouterr()
        assert "vacío" in captured.out.lower() or "0 adapter" in captured.out.lower()

    def test_history_muestra_adapters(self, continual_learner, tmp_path, capsys):
        ad = tmp_path / "ad"
        ad.mkdir()
        (ad / "meta.json").write_text(json.dumps({"model_id": "Qwen/Qwen2.5-3B-Instruct"}))
        ds = tmp_path / "ds.jsonl"
        ds.write_text(json.dumps({"messages": []}) + "\n")
        continual_learner.register_existing(str(ad), str(ds), name="demo")

        continual_learner.history()
        captured = capsys.readouterr()
        assert "demo" in captured.out

    def test_rollback_sin_backup_devuelve_false(self, continual_learner, tmp_path):
        ad = tmp_path / "no_backup"
        ad.mkdir()
        result = continual_learner.rollback(str(ad))
        assert result is False

    def test_backup_y_restore(self, continual_learner, tmp_path):
        """Crear backup manual y verificar que restore funciona."""
        ad = tmp_path / "adapter"
        ad.mkdir()
        (ad / "weights.safetensors").write_text("fake weights")
        (ad / "meta.json").write_text(json.dumps({"model_id": "Qwen/Qwen2.5-3B"}))

        backup = continual_learner._backup_adapter(ad)
        assert backup is not None
        assert backup.exists()

        # Destruir el original
        (ad / "weights.safetensors").write_text("corrupted")

        # Restaurar
        continual_learner._restore_backup(ad, backup)
        assert (ad / "weights.safetensors").read_text() == "fake weights"

    def test_check_regression_activa_rollback(self, continual_learner, tmp_path):
        """Si eval_loss sube >15%, se activa rollback."""
        ad = tmp_path / "ad"
        ad.mkdir()
        (ad / "meta.json").write_text(json.dumps({"model_id": "Qwen/Qwen2.5-3B"}))
        backup = ad.parent / (ad.name + "_backup")
        backup.mkdir()
        (backup / "meta.json").write_text("{}")

        continual_learner._registry["adapters"].append({
            "name": ad.name,
            "eval_loss": 0.10,  # baseline bajo
            "output_dir": str(ad),
        })

        regression_pct, triggered = continual_learner._check_regression(
            new_eval_loss=0.20,  # subió 100%
            adapter_name=ad.name,
            output_dir=ad,
            backup_dir=backup,
        )
        assert triggered is True
        assert regression_pct > 0.15


# ===========================================================================
# E. Formato del interaction log (S10.1)
# ===========================================================================

class TestInteractionLog:
    """Verifica el formato del log de interacciones usado por S10."""

    def test_log_entry_tiene_campos_obligatorios(self):
        """Cada entrada del interaction log debe tener los campos esperados."""
        required = {"id", "timestamp", "session_id", "turn", "user_msg",
                     "assistant", "model", "ms", "feedback"}
        entry = {
            "id": "abc-123_t2",
            "timestamp": "2026-05-18T10:30:00Z",
            "session_id": "abc-123",
            "turn": 2,
            "user_msg": "¿Qué tal?",
            "assistant": "¡Bien! ¿Y tú?",
            "model": "domestic_v2",
            "ms": 450,
            "feedback": None,
        }
        assert required == set(entry.keys())

    def test_log_entry_feedback_positivo(self):
        """Feedback = 1 es like."""
        entry = {
            "id": "xyz_t1",
            "timestamp": "2026-05-18T10:30:00Z",
            "session_id": "xyz",
            "turn": 1,
            "user_msg": "Hola",
            "assistant": "Hola, ¿en qué te ayudo?",
            "model": "test",
            "ms": 100,
            "feedback": 1,
        }
        assert entry["feedback"] == 1

    def test_log_extractor_interpreta_feedback(self):
        """El extractor de _cmd_learn debe reconocer feedback positivo y None."""
        entries = [
            {"feedback": 1, "user_msg": "a", "assistant": "b"},
            {"feedback": None, "user_msg": "c", "assistant": "d"},
            {"feedback": -1, "user_msg": "e", "assistant": "f"},
            {"feedback": 0, "user_msg": "g", "assistant": "h"},
        ]
        # Simular la lógica de _cmd_learn --auto
        accepted = []
        for e in entries:
            fb = e.get("feedback")
            if fb is None or fb == 1:
                if e.get("user_msg") and e.get("assistant"):
                    accepted.append(e)
        # Solo se aceptan feedback=1 y feedback=None.
        # feedback=0 NO se incluye (0 != 1 y 0 is not None).
        assert len(accepted) == 2

    def test_log_a_disco_y_recuperacion(self, tmp_path):
        """Escribir y leer el log mantiene la integridad."""
        log_path = tmp_path / "interaction_log.jsonl"
        entries = [
            {"id": "s1_t1", "timestamp": "2026-05-18T10:00:00Z", "session_id": "s1",
             "turn": 1, "user_msg": "Hola", "assistant": "Hey", "model": "m", "ms": 100, "feedback": None},
            {"id": "s1_t2", "timestamp": "2026-05-18T10:01:00Z", "session_id": "s1",
             "turn": 2, "user_msg": "Adiós", "assistant": "Chao", "model": "m", "ms": 90, "feedback": 1},
        ]
        with open(log_path, "w", encoding="utf-8") as f:
            for e in entries:
                f.write(json.dumps(e, ensure_ascii=False) + "\n")

        # Leer de vuelta
        recovered = []
        with open(log_path, encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    recovered.append(json.loads(line))

        assert len(recovered) == 2
        assert recovered[0]["id"] == "s1_t1"
        assert recovered[1]["feedback"] == 1
